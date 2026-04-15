from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT_JSON = ROOT / 'outputs' / 'evaluation_report.json'
GENERATED_JSON = ROOT / 'outputs' / 'generated_outputs.json'
OUT_DOCX = ROOT / 'report' / 'Final_Report.docx'


def set_base_styles(doc: Document):
    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(10.5)
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        style = styles[style_name]
        style.font.name = 'Aptos'
        style.font.color.rgb = RGBColor(31, 78, 121)
    styles['Heading 1'].font.size = Pt(18)
    styles['Heading 1'].font.bold = True
    styles['Heading 2'].font.size = Pt(13)
    styles['Heading 2'].font.bold = True
    styles['Heading 3'].font.size = Pt(11)
    styles['Heading 3'].font.bold = True


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Email Generation Assistant\nFinal Report')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('AI Engineer Candidate Assessment Deliverable').italic = True


def add_summary_box(doc: Document, summary: dict):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for c, t in zip(hdr, ['Model', 'Fact Coverage', 'Tone Alignment', 'Structure & Fluency', 'Overall']):
        c.text = t
    for model_name, stats in summary.items():
        row = table.add_row().cells
        row[0].text = model_name
        row[1].text = str(stats['avg_fact_coverage_score'])
        row[2].text = str(stats['avg_tone_alignment_score'])
        row[3].text = str(stats['avg_structure_fluency_score'])
        row[4].text = str(stats['avg_overall_score'])
    doc.add_paragraph()


def add_prompt_section(doc: Document, generated: list[dict]):
    sample = generated[0]
    doc.add_heading('1. Prompt Template Used', level=1)
    doc.add_paragraph(
        'The recommended production strategy uses advanced prompt engineering that combines role-playing, '
        'instruction scaffolding, and few-shot guidance. The baseline comparison uses a minimal prompt.'
    )
    doc.add_heading('Advanced Prompt Preview', level=2)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.add_run(sample['advanced_prompt_preview'])
    doc.add_heading('Baseline Prompt Preview', level=2)
    doc.add_paragraph(sample['baseline_prompt_preview'])


def add_metrics_section(doc: Document, metric_definitions: dict):
    doc.add_heading('2. Custom Metric Definitions and Logic', level=1)
    for i, (name, desc) in enumerate(metric_definitions.items(), start=1):
        doc.add_heading(f'Custom Metric {i}: {name}', level=2)
        doc.add_paragraph(desc)
    doc.add_paragraph(
        'These metrics are tailored to the email-generation task rather than generic text evaluation. '
        'Together, they measure factual completeness, tone consistency, and professional email quality.'
    )


def add_raw_data_section(doc: Document, rows: list[dict]):
    doc.add_heading('3. Raw Evaluation Data', level=1)
    doc.add_paragraph('The table below shows the raw scores for all 10 scenarios across both strategies.')
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Scenario', 'Model', 'Tone', 'Fact', 'Tone', 'Structure', 'Overall']
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for rowd in rows:
        row = table.add_row().cells
        row[0].text = str(rowd['scenario_id'])
        row[1].text = rowd['model']
        row[2].text = rowd['tone']
        row[3].text = str(rowd['fact_coverage_score'])
        row[4].text = str(rowd['tone_alignment_score'])
        row[5].text = str(rowd['structure_fluency_score'])
        row[6].text = str(rowd['overall_score'])
    doc.add_paragraph()


def add_analysis(doc: Document, summary: dict, rows: list[dict]):
    doc.add_heading('4. Comparative Analysis Summary', level=1)
    adv = summary['advanced_prompt']
    base = summary['baseline_prompt']
    doc.add_paragraph(
        f"The advanced prompting strategy performed better overall. It achieved an average overall score of {adv['avg_overall_score']} "
        f"compared with {base['avg_overall_score']} for the baseline strategy. The biggest advantage came from fact coverage "
        f"(1.0 vs. {base['avg_fact_coverage_score']}) and structure quality ({adv['avg_structure_fluency_score']} vs. {base['avg_structure_fluency_score']})."
    )

    baseline_rows = [r for r in rows if r['model'] == 'baseline_prompt']
    low_tone = sum(1 for r in baseline_rows if r['tone_alignment_score'] <= 0.25)
    doc.add_paragraph(
        f"The biggest failure mode of the lower-performing baseline strategy was incomplete and stylistically weak email generation. "
        f"In particular, it often dropped the third required fact and used generic openings and closings that did not match the requested tone. "
        f"This is visible in tone alignment, where {low_tone} out of 10 baseline scenarios scored 0.25 or lower."
    )
    doc.add_paragraph(
        "For production, I recommend the advanced prompt strategy. It is more reliable because it combines a clear role definition, "
        "explicit formatting rules, and a few-shot example. Those controls translate directly into better factual recall, more consistent "
        "professional structure, and better tone handling."
    )


def add_appendix(doc: Document, generated: list[dict]):
    doc.add_heading('Appendix A. Scenario Inventory', level=1)
    for item in generated:
        doc.add_heading(f"Scenario {item['id']}: {item['intent']}", level=2)
        doc.add_paragraph(f"Tone: {item['tone']}")
        for fact in item['key_facts']:
            doc.add_paragraph(fact, style='List Bullet')


def build():
    report = json.loads(REPORT_JSON.read_text())
    generated = json.loads(GENERATED_JSON.read_text())
    doc = Document()
    set_base_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_title(doc)
    doc.add_paragraph()
    add_summary_box(doc, report['summary'])
    add_prompt_section(doc, generated)
    add_metrics_section(doc, report['metric_definitions'])
    add_raw_data_section(doc, report['scenario_scores'])
    add_analysis(doc, report['summary'], report['scenario_scores'])
    doc.add_page_break()
    add_appendix(doc, generated)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Wrote {OUT_DOCX}')


if __name__ == '__main__':
    build()
